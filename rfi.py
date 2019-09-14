# Peter Sterckx
# Spring 2019
# Generates a Word Document of the transcript of the daily "Journal en francais facile" from RFI using BeautifulSoup


from urllib.request import urlopen
from bs4 import BeautifulSoup
import smtplib
import datetime
from docx import Document
from docx.shared import Pt, Length

failed = True
date_q = input('Please input a date (mm-dd-yyyy) or type \"today\": ')

while(failed):

    if date_q == 'today':

        date = datetime.datetime.today()

        if date.month < 10:
            date_id = str(date.day) + '0' + str(date.month) + str(date.year)
        else:
            date_id = str(date.day) + str(date.month) + str(date.year)
    else:

        date_id = date_q[3:5] + date_q[0:2] + date_q[6:10]

    url = 'https://savoirs.rfi.fr/fr/apprendre-enseigner/langue-francaise/journal-en-francais-facile-'+ date_id + '-20h00-gmt'

    try:
        uClient = urlopen(url)
        failed = False
    except:
        date_q = input('RFI not available for date specified. Please input a new date (mm-dd-yyyy): ')

page = uClient.read()
uClient.close()

soup = BeautifulSoup(page,"html.parser")

selection = soup.find_all("div", class_="field field-name-field-descriptif-trad field-type-text-long field-label-hidden")

text = selection[0].get_text()

document = Document()

style = document.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

paragraph = document.add_paragraph('RFI ' + date_q[0:2] + '-' + date_q[3:5] + '-' + date_q[6:10])
paragraph = document.add_paragraph(text)
paragraph.style = document.styles['Normal']
paragraph_format = document.styles['Normal'].paragraph_format
paragraph_format.line_spacing = Pt(14)

filename = 'RFI' + str(date_q[0:2] + date_q[3:5] + date_q[6:10])
document.save(filename + '.docx')

print("Transcript created.")
